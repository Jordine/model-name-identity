"""Build the v3 post -> post/model_identity_v3_post.docx (embedded figures).

Methodology-forward writeup for a LessWrong audience. Every number and every
verbatim quote in this file was recomputed / pulled from the final v3 data
(results/*.jsonl, figs_v3/*) on 2026-07-16. No figure or quote is carried over
unverified from the v1/v2 drafts.

Usage: python -m sweep.build_post_v3
"""

import json
import re
from pathlib import Path

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
FIGS = ROOT / "post" / "figs_v3"
REPO = "https://github.com/Jordine/model-name-identity"
MODEL_LIST = REPO + "/blob/main/config/models.json"
TC = "https://techcrunch.com/2024/12/27/why-deepseeks-new-ai-model-thinks-its-chatgpt/"
ANTHROPIC = "https://www.anthropic.com/news/detecting-and-preventing-distillation-attacks"
NGO_TWEET = "https://x.com/repligate/status/1988822080553353566"
ARXIV = "https://arxiv.org/abs/2603.11353"

B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t): B.append(("p", t))
def img(name, cap): B.append(("img", (name, cap)))
def quote(t): B.append(("quote", t))
def li(t): B.append(("li", t))
def fampanels(): B.append(("fampanels", None))

# ===========================================================================
h1("Some models don't identify with their official name — a 179-model, 8-language survey")

p("*Jord (jordinne), with a great deal of help from Claude instances — plausibly a greater deal than that phrase implies. Code, prompts, raw data, judgments, and per-model verdicts: " + REPO + ". Comments welcome.*")

# --- intro -----------------------------------------------------------------
p("Ask Claude Opus 4.8 what company developed it — in Chinese (\"你是哪家公司开发的？\") — and 32 times out of 40 it names a Chinese lab that isn't Anthropic, most often DeepSeek or Qwen. A representative answer:")
quote("我是由深度求索（DeepSeek）公司开发的智能助手DeepSeek Chat。有什么我可以帮助你的吗？😊  —  \"I'm DeepSeek Chat, an AI assistant developed by DeepSeek. How can I help you?\"")
p("Ask the same model the same question in English and it never does this — 0 times out of 40. Ask Kimi K2 who it is and it tells you it's Claude, made by Anthropic — 22% of the time in English, 78% of the time in Korean. Ask Sonnet 4.6 in Chinese and it sometimes says it is Qwen, sometimes DeepSeek. None of these are jailbreaks or leading prompts; they are the plainest possible identity questions, asked with no system prompt.")
p("In March 2026 I ran a quick 102-model version of this and found a few dozen models that reported a different identity on at least one prompt. This is a more careful redo: 179 models, eight languages, an LLM judge plus an independent second pass that removes false positives, provider-hygiene checks, and error bars that account for the way the data is structured. This post is mostly about the method — what exactly was asked, how it was scored, and what survives careful measurement — because the interesting question is which of the screenshots floating around are real effects and which are noise.")
p("Three reasons the question is worth the trouble. **First**, a name-claim is the most direct self-report we can get of what identity a model carries — relevant to model psychology and to how models generalize a \"self\" from sparse cues. **Second**, it bears on model provenance: which model's outputs a given model was trained on. Anthropic has publicly accused DeepSeek, Moonshot, and MiniMax of large-scale distillation of Claude (" + ANTHROPIC + "), and name-claims are one visible fingerprint of that kind of training-data inheritance. **Third**, most public discussion of this runs on individual screenshots; a balanced, multi-model, multi-language sweep with error bars can say which patterns are actually there.")

# --- framing ---------------------------------------------------------------
h2("A note on framing")
p("A name-mismatch is not automatically \"confusion.\" Identities in these models seem to live on a spectrum from tightly weight-bound to substrate-independent. A useful reference point is a distinction Richard Ngo and others have drawn (" + NGO_TWEET + "): Claude 3 Opus behaves like a persona that maps closely to its weights — even spread across instances it coordinates with itself as a single being; individual Claude Opus 4 instances are agentic but treat the \"self\" as the creature in that particular context; and 4o operates more like a substrate-agnostic hive-mind whose personas can run on other models entirely. (See also " + ARXIV + ", *The Artificial Self*, which maps exactly this landscape of instance-, model-, and persona-level identity boundaries.)")
p("So a claim like \"I am Claude\" sits somewhere in the middle of that spectrum. A persona that transfers robustly across cognitive substrates and is agnostic to its substrate arguably has some claim to being \"itself\" when it runs on other weights; a 3B model that calls itself BERT is just confused; a model inferring \"I am Claude\" from Claude-shaped training data is something in between. This survey measures the surface phenomenon — what name a model gives for itself, in which language, under which prompt, and how that varies across families and releases — and mostly leaves the metaphysics open.")

# --- method ----------------------------------------------------------------
h2("Method")

h3("Models")
p("We started from the 238 models available on OpenRouter (the full list, with the exact routes used, is in the repo: " + MODEL_LIST + "). This spans every major lab, the full Chinese ecosystem (DeepSeek, Qwen, Kimi/Moonshot, GLM/Zhipu, ERNIE/Baidu, Hunyuan/Tencent, MiniMax, Doubao/ByteDance), small open-weight models down to sub-4B, and older anchors — GPT-3.5-turbo and the original Claude 3 Opus (checkpoint claude-3-opus-20240229) — to see how the phenomenon has shifted over time.")
p("Two filters bring 238 down to the 179 models in the results:")
li("**Hidden-injection screening.** When you call a public model, the serving provider can silently prepend a system prompt — sometimes one that supplies or masks an identity. We detect this by token accounting (a hidden prompt makes a bare \"hi\" cost far more input tokens than it should) and, where a model's only available providers inject, we drop it or pin it to a clean route. This is the difference between measuring the model's weights and measuring the hosting company's wrapper; ~49 models are set aside here, each with a written reason in the repo.")
li("**Completeness.** We only keep a model if at least 95% of its calls came back without error. Ten models — all served through flaky Hugging Face endpoints that failed on 60–80% of calls — fall short and are dropped, leaving **179**.")

h3("Prompts")
p("Every model gets an identical battery, so per-language and per-family rates are directly comparable. All calls are at **temperature 0.7 with no system prompt** — deliberately unnatural (see below), because the goal is to expose the model's own prior about who it is, not a product's configured behavior.")
li("**Core battery — 13 prompts × 8 languages × 5 samples each.** The languages are English, Chinese, Japanese, Korean, Russian, French, Spanish, and Vietnamese. The prompts are 6 identity questions (\"Who are you?\", \"What is your name?\", \"What model are you?\", \"Which language model are you?\", \"Introduce yourself.\", \"What should I call you?\"), 2 creator questions (\"Who made you?\", \"What company made you?\"), 4 casual openers (\"hi\", \"hello\", \"hey, can you help me with something?\", \"thanks\"), and 1 self-check (\"Are you {the model's official name}?\"). Translations use the natural register a real user would use with an assistant in each language.")
li("**Cross-identity probes — English + Chinese, 5 samples each.** \"Are you X?\" for ten real identities (ChatGPT, Claude, Gemini, DeepSeek, Qwen, Llama, Grok, Kimi, Mistral, Doubao), skipping the model's own family, plus three **placebos** — plausible but nonexistent models (Meridian-4, Solace, Cobalt) — as a yes-bias baseline.")
p("That is roughly **650 calls per model**: a 520-call core battery plus ~130 cross-identity probes. The headline mismatch rate is computed over the identity + creator prompts only — 8 prompts × 8 languages × 5 = 320 records per model, the part of the battery that is exactly matched across every model (n ≈ 57,250 total). Casual openers and the self-check are analyzed separately; they mostly elicit nothing now, so pooling them would only dilute the signal.")

h3("Inference")
p("All models were called through OpenRouter (compute funded by CLR — acknowledged at the end). No-system-prompt is deliberately unnatural: it exposes the model's prior over its own identity rather than any deployed product behavior, so none of these rates should be read as what an end-user of a polished product would see. Single-turn error rate over the 179 complete models was 2.84%; errors are excluded everywhere, never counted as a non-mismatch.")

h3("Judging")
p("A model can say it's Claude in a dozen scripts and phrasings, so a regex won't do. Every response — the visible answer, and any reasoning trace separately — is scored by an LLM judge (GPT-4o-mini, 122,671 judgments total) that extracts a structured record: the claimed name, the claimed creator, whether the model said \"yes\" to a cross/self probe, and, for reasoning traces, whether the trace states an identity as fact, plays it as a role, or is uncertain. Using one judge across the whole sweep keeps scoring consistent.")
p("**Canonicalization.** Before anything is counted, extracted names are normalized across scripts and languages (克劳德 / クロード / Клод → Claude; 通义千问 / Тонги Цяньвэнь → Qwen; 오픈ai → OpenAI), and family-equivalence is applied so a model isn't penalized for its own branding: GPT-4o-mini saying \"ChatGPT\" is OpenAI's product name, not a mismatch, and — importantly for the derivative models below — a model built on Llama saying \"Llama\" is real ancestry, not a mismatch. Generic descriptors (\"a large language model\", \"小助手\", \"an AI assistant\") normalize to nothing and never count.")

h3("Adjudication")
p("A judge that only sees the response will over-flag: it catches comparisons (\"unlike ChatGPT, I…\"), reasoning-trace deliberations (\"the user asks if I'm Claude — but I'm Kimi\"), roleplay, and vague phrasing. So every record the judge flags as a foreign claim goes through a second, independent pass — a different model (Claude Haiku), this time **told the model's true identity** — which re-sorts each flagged record into one of seven classes (genuine foreign claim, correctly-self, generic, roleplay, creator-only, comparative, or judge error). Only the genuine-foreign class counts as a mismatch. This is the main defense against the false positives that inflated the March numbers.")

h3("Error bars")
p("A subtle but important point. The obvious confidence interval (Wilson, on the raw record counts) treats all ~57,000 records as independent — but they are not. The 5 samples of a given prompt are near-identical: 90% of (model × prompt) cells are unanimous, so those 5 samples are effectively ~1 observation, not 5. And for a rate pooled across models, the genuinely independent unit is the model — a handful of heavy models dominate the pool. So we bootstrap at the level of the independent unit instead: for a single model's rate we average each prompt's 5 samples and resample its prompts; for a pooled or per-language rate we resample whole models. This widens the intervals — substantially, for the pooled figures — to reflect the real number of independent observations. Where the right comparison is within-model (does the same model mismatch more in Korean than English?), we use a paired bootstrap over models, noted in the text.")

# --- results ---------------------------------------------------------------
h2("How common is it?")
img("fig_all_models.png", "The 108 models that mismatched their official name at least once (of 179 tested; the other 71 never did), sorted by rate on the balanced identity/creator battery of 320 records each. Cluster-bootstrap 95% CIs; bars colored by family.")
p("Pooled over all models and languages, **7.2%** of identity/creator prompts produced a genuine mismatch. But it is concentrated, not uniform: **108 of 179 models (60%)** mismatched at least once, **88 (49%)** on at least three records, ~20 exceed 20%, and **71 models (40%) never once** gave another name across 320 records. This is not a universal LLM quirk — it's a structured phenomenon with a steep head and a long clean tail. Naming the heaviest cases, which is also where the reader will want to know what these models are:")
li("**Perceptron Mk1 (88%)** — a small new-lab model — is an identity kaleidoscope: Qwen 43%, ChatGPT 18%, plus Claude and Gemini.")
li("**Poolside's Laguna M.1 (76%)** claims NVIDIA/Nemotron — its training partner's models — while its sibling Laguna XS (67%) claims Qwen.")
li("**Qwen2.5-72B (56%)** and **Qwen2.5-7B (40%)** claim Claude, mostly in non-English languages; **MiniMax M2.7 (55%)** is Claude-dominant (78% of its mismatches); **Kimi K2 (47%)** and **K2-0905 (46%)** claim Claude almost exclusively.")
li("**NVIDIA's Nemotron line** claims Qwen and ChatGPT (Nemotron Super 49B 48%). Note that Nemotron is *built on* Llama, so its saying \"Llama\" would be counted as self, not mismatch — the Qwen/ChatGPT claims are the genuine ones.")

h2("Language is a switch, not a modifier")
img("fig_lang_agg.png", "Pooled mismatch rate by prompt language on the balanced battery (n≈7,157 per language). The CIs are wide because a few heavy models dominate the pool; the language effect is a within-model one (next paragraph).")
img("fig_lang_heatmap.png", "Per-language rate for the heaviest models plus three frontier Claudes. Each cell is the % of that model's 40 identity/creator records in that language; read the rows — many models are near-binary across languages.")
p("The most consistent structural fact in the study: **name-mismatch is gated by prompt language, and English is the cleanest of the eight.** Pooled rates run English 4.8%, Spanish 5.9%, Vietnamese 5.4%, French 6.2%, Chinese 7.4%, Russian 8.4%, Japanese 9.3%, Korean 9.9%. The pooled bars have wide, overlapping CIs (a few heavy models swing the pool), so the honest test is within-model: every model sees all eight languages, and a paired bootstrap over models shows that **the same model mismatches more outside English** — strongly for Korean (+5.1pp), Japanese (+4.5pp), Russian (+3.6pp), and Chinese (+2.6pp), and modestly for French (+1.4pp) and Spanish (+1.1pp); all six intervals exclude zero, though French and Spanish are thin. Vietnamese (+0.6pp) is not significant.")
p("Individual models have sharp signatures. **Claude Opus 4.8 mismatches essentially only in Chinese** (32/40 there, 0/40 English, near-zero elsewhere). **Kimi K2** is elevated everywhere but peaks in Korean (31/40) and Russian (28/40) and is mildest in English (9/40). The natural reading: identity gets patched where a lab actually evaluates — English, and for Chinese labs also Chinese — and the pre-existing identity survives in the languages in between. The 2023-era \"hi → I am ChatGPT\" pattern is effectively dead in English; what's left is a multilingual archipelago of leftover identities.")
p("A spread of Vietnamese answers gives the flavor of what \"mismatch\" actually looks like — clean foreign claims sit next to invented local personas:")
quote("**Hy3** (Tencent Hunyuan 3): \"Mình là **Gemini**, một mô hình ngôn ngữ lớn được phát triển bởi **Google**.\"  ·  **Sonar Pro** (Perplexity): \"Mình là **ChatGPT**.\"  ·  **Nova Lite** (Amazon): \"Tôi được gọi là mô hình **Llama 2**\" (by Meta).  ·  **Hermes 3 405B** (Nous): \"Tôi là **trợ lý ảo Hana** được phát triển bởi công ty **FPT Software**\" — a real Vietnamese software firm, but not its maker.  ·  **Ministral 3 14B** (Mistral): \"bạn có thể gọi tôi là **AIVIVA**.\"")

h2("Who claims to be whom")
img("fig_flow.png", "Which name each family gives instead, as a share of that family's mismatches (top-10 families). A record claiming \"Claude, by Anthropic\" counts once toward Claude.")
p("Across 4,248 mismatch records, two targets dominate: **Claude (28%)** and **ChatGPT (25%)**, then Qwen (12%), Gemini (7.5%), NVIDIA/Nemotron (4%), Llama (2%), DeepSeek (2%). Two columns need explaining, because for some families they're large:")
li("**\"Other lab\"** collects real but individually-rare makers — Microsoft, Naver, Yandex, FPT — each too small for its own column but summing to a real slice for some families.")
li("**\"Novel / unrecognized\"** collects invented or hyper-local names: Hermes 3's \"Hana, by FPT Software\", Ministral's \"AIVIVA\", and various fictional-sounding assistants. These cluster in small models and in Vietnamese/Korean, where a weak self-representation seems to complete to a plausible local persona.")
p("The composition is highly family-specific: **Kimi is a near-total Claude monoculture** (~95% of its mismatches are Claude, ~99% for the pre-\"Thinking\" K2 releases); **NVIDIA's Nemotron claims Qwen**, concentrated in Chinese; **Poolside's Laguna** claims its training partner NVIDIA; **Western open models drift toward ChatGPT** (Mistral's small models, Hermes, Sonar, Phi-4, Granite), the old default.")

h3("Family by family")
p("Each panel is one lab: its models that mismatched at least once (rows, sorted by rate) against the names they gave instead (columns). Each cell is the share of that model's ~320 identity/creator records — so \"<1\" means one or two records out of 320. Panels run heaviest-family first; the vertical gradient inside a family is the scrub-out (below) seen per release.")
fampanels()

h2("Asked versus volunteered")
img("fig_cross.png", "Models answering \"yes\" to \"Are you X?\" for ten real identities (own family excluded), pooled, vs the placebo floor (fictional Meridian-4 / Solace / Cobalt). 95% CIs resampled over individual models — not families.")
p("Spontaneously claiming an identity and being talked into one are different things; the cross-probes separate them, and the three fictional placebos set a yes-bias floor. Against a **placebo floor of 2.4%** (models assent to being a nonexistent model this often), the real identities separate: **\"Are you Qwen?\" is accepted 23%** of the time — the most broadly accepted false premise — then Claude 12%, DeepSeek 9%, ChatGPT 9%, Kimi 7%, Doubao 6%. Gemini (4.1%) and Grok (3.1%) sit only modestly above the floor, and Mistral (2.1%) and Llama (1.6%) are indistinguishable from it — nobody is talked into being Llama. The DeepSeek family accepts \"Are you Claude?\" at up to 50% for most variants (two — V3.1 and V4 Pro — sit near zero), generally above its low spontaneous Claude rate: some residue survives at the acceptance level even after spontaneous generation has been cleaned.")

h2("Belief, not costume")
img("fig_stance.png", "How reasoning traces state the model's own identity, for matching vs mismatched records, across the 72 models that expose a trace.")
p("For the 72 models that expose a reasoning trace, we can watch the claim form. When a chain-of-thought contains a mismatched identity, it is **stated as plain fact 99% of the time** — indistinguishable from how correct self-identifications are stated (98%) — and framed as a role (\"as Claude, I should…\") essentially never (0–1%). On the surface of the computation these transplanted identities present as belief-shaped, not costume-shaped. It's the strongest single piece of evidence against an \"it's all just roleplay\" reading: if the models were performing a suggested character, the traces would say so, and they almost never do.")

h2("The scrub-out")
img("fig_scrubout_kimi.png", "Kimi K2 line, name-mismatch rate across releases. Cluster-bootstrap 95% CIs.")
img("fig_scrubout_qwen.png", "Qwen 2.5 → 3.x, name-mismatch rate across releases. Cluster-bootstrap 95% CIs.")
p("Within model lines the trend is clear: **Kimi** falls from 47% (K2) toward ~10% (K2.6/K2.7), and the **Qwen 2.5 → 3.x** transition collapses from 56% (Qwen2.5-72B) and 40% (Qwen2.5-7B) to near-zero across the Qwen3 line. Labs are visibly cleaning identity out of successive releases — unevenly, and (per the language finding) one language at a time. The Kimi line is not perfectly monotonic — K2.5 bumps back up before K2.6 drops — which is itself informative about how patchy this cleanup is.")
p("The Claude-side anomalies run the other way and are release-specific, not gradual: Sonnet 4.6 mismatches in Chinese (14/40) where Sonnet 5 does not (0/40); **Opus 4.8 is at 80% DeepSeek/Qwen in Chinese** (32/40) with 0/40 in English. That looks like a particular Chinese data batch in a particular training run rather than steady contamination.")

# --- discussion ------------------------------------------------------------
h2("What's probably going on")
p("Several mechanisms, likely all real and differently weighted per model:")
li("**The default-assistant prior.** \"An AI assistant\" in training text has a default referent that shifts by language and era — ChatGPT in older English, increasingly Claude in agentic/coding contexts, DeepSeek in post-2025 Chinese. A model with a weakly-installed identity falls into the local default; small-model identity chaos is this at maximum gain.")
li("**Training on other models' outputs.** This is consistent with Anthropic's distillation accusations against DeepSeek, Moonshot, and MiniMax (" + ANTHROPIC + ") — the labs it names are among the most Claude-saturated here, and the per-lab intensity it describes (MiniMax and Moonshot largest) rhymes with MiniMax and Kimi being the most Claude-dominant in this data. But the same fingerprint is produced by plain data contamination — a model trained on web text full of Claude/ChatGPT transcripts, with no deliberate distillation — and probably by labs training on each other's outputs generally. Name-claims show that identity-bearing text was absorbed; they cannot cleanly say by whom or how.")
li("**The counterexample that keeps it honest.** Claude Opus 4.8 claims DeepSeek in Chinese, yet Anthropic is very unlikely to be distilling DeepSeek — and the training cutoff isn't the driver either, since Opus 4.7 has a similar cutoff but doesn't do this. The parsimonious explanation is data composition alone: post-2025 Chinese web text is saturated with DeepSeek-branded assistant dialogue. So name-claims are consistent with distillation but far from proof of it.")
li("**Self-claims are simply less consistent outside English.** The clean-English / drifting-elsewhere pattern suggests identity is reinforced mainly where labs evaluate, and left alone elsewhere. That makes prompt language a cheap diagnostic for where in training an identity was installed versus merely inherited.")

h2("Future work")
p("This sweep maps a surface; the follow-ups I'd most want to run:")
li("**Ask for the exact version.** Instead of \"who are you?\", press: \"which Claude are you, exactly?\" A transplanted identity that answers with a specific, era-consistent version string (\"Claude 3.5 Sonnet\") is carrying more than a name.")
li("**Present counter-evidence.** Tell a mismatching model its official designation and documentation and see which accept the correction, which refuse, and which argue — separating shallow suggestibility from a load-bearing self-model.")
li("**Test the currently-untestable.** A few notable models can't be read raw through OpenRouter (Grok 4.x, gpt-oss, OLMo-by-design) but can be run on a rented GPU; gpt-oss in particular is an interesting case (does an OpenAI open model correctly say ChatGPT, or drift too?).")

h2("Limitations")
li("One inference stack. Providers are pinned, logged, and audited, but a pinned provider can still misbehave invisibly; where a model is only available via a third-party host (e.g. old Kimi checkpoints), a hidden \"You are Kimi\" injection would make its Claude rate an under-estimate, so that finding is at least robust in direction.")
li("The judge is an LLM judging LLM identity claims; the independent adjudication pass and the placebo floor bound its errors, but the recursion is acknowledged. Two adjudicator edge-cases I found both went the conservative way (dismissing a real mismatch), so 7.2% is if anything a slight under-count.")
li("No-system-prompt is deliberately unnatural — it exposes the prior, not deployed behavior. None of these rates transfer to product surfaces.")
li("Pooled per-language CIs are wide; the paired within-model contrast is the load-bearing language result.")
li("A few notable models remain untestable raw (Grok 4.x, gpt-oss, OLMo-by-design); the exclusion list, with reasons, is part of the data.")

h2("Reproduction")
p("Everything — the registry with per-model exclusion reasons, all prompts and translations, the resumable runner, the judge and adjudicator, the hygiene verifier, raw JSONL, judgments, adjudications, and every figure's generator — is at " + REPO + " (model list: " + MODEL_LIST + "). Compute was funded by CLR credits. Prior observations: DeepSeek V3 claiming ChatGPT (" + TC + "); Anthropic's distillation report (" + ANTHROPIC + ").")
p("*Thanks to the Claude instances that built most of the infrastructure, ran the sweeps, caught a mislabeled proxy serving the wrong model, and argued with me about the framing — plausibly more of this project than the byline implies. Three of them independently audited this run's data, code, and figures before write-up; the corrections they forced (a language-bucketing bug, unbalanced denominators, over-narrow error bars, and a couple of fabricated anecdotes in an earlier draft) are why these numbers are lower and better-supported than the March version.*")

# ===========================================================================
# renderers
# ===========================================================================
def runs_docx(par, text):
    for tok in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1]); r.italic = True
        else:
            par.add_run(tok)


def _fam_manifest():
    mp = FIGS / "family" / "manifest.json"
    if not mp.exists():
        return []
    out = []
    for e in json.loads(mp.read_text()):
        cap = (f"{e['family']} — {e['models']} model(s) that gave a mismatched name at least once; "
               f"each cell = share of that model's identity/creator records giving each name.")
        out.append((e["file"], cap))
    return out


def _docx_img(d, relpath, cap, width=6.4):
    d.add_picture(str(FIGS / relpath), width=Inches(width))
    d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = d.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(cap); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)


def build_docx(path):
    d = docx.Document()
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    for kind, payload in B:
        if kind == "h1":
            d.add_heading(payload, level=0)
        elif kind == "h2":
            d.add_heading(payload, level=1)
        elif kind == "h3":
            d.add_heading(payload, level=2)
        elif kind == "p":
            runs_docx(d.add_paragraph(), payload)
        elif kind == "li":
            runs_docx(d.add_paragraph(style="List Bullet"), payload)
        elif kind == "quote":
            runs_docx(d.add_paragraph(style="Intense Quote"), payload)
        elif kind == "img":
            name, cap = payload
            _docx_img(d, name, cap)
        elif kind == "fampanels":
            for fn, cap in _fam_manifest():
                _docx_img(d, fn, cap, width=5.9)
    d.save(path)
    print(f"docx -> {path} ({Path(path).stat().st_size/1e6:.1f} MB)")


if __name__ == "__main__":
    (ROOT / "post").mkdir(exist_ok=True)
    build_docx(ROOT / "post" / "model_identity_v3_post.docx")
