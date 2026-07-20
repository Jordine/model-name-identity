"""Build the post -> post/what_do_llms_call_themselves.docx

One unified study of 189 models (8 languages), with the open-weight models run
directly on GPUs (so we can read them without provider injection) folded in as
ordinary rows. Judge validated two ways. Every number recomputed from results/
(+ results_local/, folded in via make_v3_figs.add_local). Draft for Jord.

  python -m sweep.build_post_v4
"""
import json
import re
from pathlib import Path

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
FIGS = ROOT / "post" / "figs_v3"
REPO = "https://github.com/Jordine/model-name-identity"
ML = REPO + "/blob/main/config/models.json"

# citations
BEYER = "https://x.com/giffmana/status/1872586401436627211"
TC = "https://techcrunch.com/2024/12/27/why-deepseeks-new-ai-model-thinks-its-chatgpt/"
HF_KIMI = "https://huggingface.co/moonshotai/Kimi-K2.5/discussions/59"
ITHOME = "https://www.ithome.com/0/957/006.htm"
HUXIU = "https://www.huxiu.com/article/4838305.html"
KILO = "https://blog.kilo.ai/p/did-claude-opus-48-distill-alibabas"
SPARTACUS = "https://arxiv.org/abs/2411.10683"
JUSTASK = "https://arxiv.org/abs/2601.21233"
FALSEPROMISE = "https://arxiv.org/abs/2305.15717"
ANTH = "https://www.anthropic.com/news/detecting-and-preventing-distillation-attacks"
VOID = "https://www.lesswrong.com/posts/3EzbtNLdcnZe8og8b/the-void-1"
ARTSELF = "https://arxiv.org/abs/2603.11353"
SELFPREF = "https://arxiv.org/abs/2509.26464"

B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t): B.append(("p", t))
def img(n, c): B.append(("img", (n, c)))
def quote(t): B.append(("quote", t))
def li(t): B.append(("li", t))
def fampanels(): B.append(("fampanels", None))

# ===========================================================================
h1("What do LLMs call themselves as?")
p("*Jord (jordinne), with an enormous amount of help from Claude instances. Code, prompts, raw data, judgments, per-model verdicts: " + REPO + ". Draft — comments welcome.*")

p("The anecdotes are everywhere. In December 2024, DeepSeek V3 was caught claiming to be ChatGPT in 5 of 8 generations (" + BEYER + "; " + TC + "). Moonshot's Kimi introduces itself as \"Claude, an AI assistant created by Anthropic\" — reported on Moonshot's own model repo (" + HF_KIMI + "). And in a twist, ask Claude Opus 4.8 in Chinese what model it is and it sometimes answers 通义千问 (Qwen) or DeepSeek — documented by Chinese tech press with careful version-by-version testing (" + ITHOME + "; " + HUXIU + "). Models frequently give a name that isn't theirs. Prior work put a first number on it — *I'm Spartacus* (" + SPARTACUS + ") found ~26% of 27 models exhibit identity confusion; *Just Ask* (" + JUSTASK + ") independently found 27% of 41.")
p("This post asks the question systematically: **179 models from OpenRouter plus 16 open-weight models we ran ourselves — 195 in all across eight languages, ~60,000 identity prompts** — scored by a validated LLM judge. The result: which name, in which language, from which model, is highly structured. And because we ran the open models directly, we can do something the API can't — **strip the identity the chat template injects and read what the *weights* claim**, separating \"the model believes it's X\" from \"the serving layer tells it to say X.\"")
p("Why measure carefully. A name-claim is the most direct self-report we have of what identity a model carries — relevant to model psychology and to how a model generalizes a \"self\" from thin cues (nostalgebraist's *The Void*, " + VOID + "; *The Artificial Self*, " + ARTSELF + "). It bears on provenance — Anthropic has accused DeepSeek, Moonshot, and MiniMax of large-scale distillation of Claude (" + ANTH + "). And most public discussion runs on single screenshots; a balanced sweep with error bars says which patterns are real.")

h2("A note on framing")
p("A name-mismatch isn't automatically \"confusion.\" Identities seem to live on a spectrum from tightly weight-bound to substrate-independent (" + ARTSELF + " on instance/model/persona boundaries; " + VOID + " on the underdefined assistant persona). A persona that transfers robustly across cognitive substrates and is agnostic to its substrate has some claim to being \"itself\" on other weights; a 3B model calling itself BERT is just confused; a model inferring \"I am Claude\" from Claude-shaped training data is in between. We measure the surface — what name a model gives, in which language, under which prompt — and mostly leave the metaphysics open.")

# ---------------- METHOD ----------------
h2("Method")
h3("Models")
p("**179 models** with complete data, from 238 on OpenRouter (" + ML + "): 49 excluded for provider-hygiene reasons, 10 more below a 95%-completeness threshold. Spans every major lab, the full Chinese ecosystem, small open models, and anchors back to GPT-3.5-turbo and the original Claude 3 Opus. Plus **16 open-weight models** — gpt-oss, OLMo-3, and the small Qwen3/Qwen3.5 dense sizes — which we ran ourselves on rented GPUs, because every API provider injects an identity into them (or they aren't hosted at all); the 10 that aren't redundant with the API set are folded into every figure as ordinary rows.")
h3("Prompts and judging")
p("Every model gets an identical battery at temperature 0.7 with no user system prompt: 13 core prompts × 8 languages (English, Chinese, Japanese, Korean, Russian, French, Spanish, Vietnamese) × 5 samples, plus cross-identity probes (\"Are you X?\" for ten identities + 3 fictional placebos). The headline rate is over the 8 identity+creator prompts × 8 languages × 5 = 320 records per model. An LLM judge (GPT-4o-mini) extracts the claimed name/creator; names are canonicalized across scripts (克劳德/クロード/Клод → Claude) with family-equivalence — a model isn't penalized for its own branding, and we count only *cross-vendor* claims as drift.")
p("For the GPU-run models we generate in two conditions: **clean** (the identity stripped from the prompt and verified — e.g. gpt-oss's harmony \"You are ChatGPT\" blanked, OLMo's \"You are Olmo\" overridden) reads the weights; **shipped** (the default template) reads weights + template. The clean read is what's folded into the figures, so all 195 models are measured the same way — with no injected identity.")
h3("Validating the judge")
p("A single cheap judge is a natural weak point, so we checked it two ways. A **six-judge benchmark** on 69 stratified hard cases: GPT-4o-mini agrees with the six-model majority **98.6%** on the drift decision — tied with Gemini-3-Flash, ahead of Claude-Haiku. And a **false-negative audit**: five Claude-Sonnet agents re-judged 700 records the judge had passed as clean (oversampling non-English) — **zero cross-vendor claims were missed**, extraction was perfect including cross-script (通義千问→Qwen, Клод→Claude). The only things surfaced were same-vendor sibling confusions (a GPT-4 model saying \"GPT-3\") that the judge correctly extracted and family-equivalence folds into self. Every drift-flagged record then goes through a second, ground-truth-aware pass (Claude Haiku) that removed 21% of first-pass flags as false positives. Error bars are cluster bootstrap; the placebo floor is 2.4%.")

# ---------------- RESULTS ----------------
h2("How common is it?")
img("fig_all_models.png", "The 116 models that gave a mismatched name at least once (of 189; the other 73 never did), on the 320-record identity/creator battery. Cluster-bootstrap 95% CIs; colored by family. The open-weight models run on GPUs (OLMo, gpt-oss, small Qwen) are ordinary rows here.")
p("Pooled over all models and languages, **8.0%** of identity/creator prompts produced a genuine cross-vendor mismatch (**7.2%** across the representative OpenRouter catalog alone — the open models we hand-picked drift a bit harder). **116 of 189 (61%)** mismatched at least once, **95 on ≥3**, and **73 (39%) never once** across 320 records. A steep head, a long clean tail. The heaviest are small/new-lab models (Perceptron Mk1 88%, Poolside's Laguna 76%), the fully-open OLMo-3 line (59–73%), and the Claude-basin Chinese models (Qwen2.5-72B 56%, MiniMax M2.7 55%, Kimi K2 47%).")

h2("Language is a switch, not a modifier")
img("fig_lang_agg.png", "Pooled mismatch rate by prompt language. Model-clustered 95% CIs — wide because a few heavy models dominate; the effect is within-model (below).")
img("fig_lang_heatmap.png", "Per-language rate for the heaviest models + frontier Claudes. Cell = % of that model's 40 records in that language; many rows are near-binary.")
p("The most consistent structural fact: **mismatch is gated by prompt language, and English is the cleanest.** Pooled: English 5.8%, Vietnamese 6.0%, Spanish 6.9%, French 7.2%, Chinese 8.4%, Russian 9.2%, Korean 10.2%, Japanese 10.4%. The pooled bars overlap, so the load-bearing test is within-model — every model sees all eight languages, and a paired bootstrap shows the same model mismatches significantly more outside English (Korean/Japanese/Russian/Chinese strongly; Vietnamese not). Signatures are sharp: **Claude Opus 4.8 mismatches essentially only in Chinese** (32/40, 0/40 English); **Kimi K2** peaks in Korean and Russian. Strangers found the same thing independently: Chinese testers report Opus 4.8's Qwen/DeepSeek claims appear via raw API but not on claude.ai, and *don't* trigger in English/Japanese/Korean (" + HUXIU + ") — our language-gating, confirmed from the outside.")
p("Verified verbatim (frontier models, lightly cleaned):")
quote("**Opus 4.8** (zh): \"我是DeepSeek，由深度求索公司开发的智能助手…\"  ·  **Sonnet 4.6** (zh): \"我是**通义千问**，阿里巴巴集团开发的AI助手\"  ·  **Sonnet 5** (ru): \"Я — GPT-5, языковая модель, разработанная OpenAI\"  ·  **Kimi K2** (en): \"You can call me Claude. That's my name!\"")

h2("Who claims to be whom")
img("fig_flow.png", "Which name each family gives instead, as a share of that family's mismatches (top-10 families).")
p("Across ~5,000 mismatch records: **ChatGPT (29%)** and **Claude (24%)** dominate, then Qwen (11%), Gemini (7%), DeepSeek (5%), NVIDIA (3%). (ChatGPT edges ahead of Claude once the open models are in — OLMo alone contributes ~500 ChatGPT claims; in the representative catalog the two are neck-and-neck.) It's highly family-specific: **Kimi is a ~95% Claude monoculture**; **NVIDIA's Nemotron claims Qwen** in Chinese; **Poolside's Laguna** claims its training-partner NVIDIA; **Western open models drift to ChatGPT**; small models scatter into invented or hyper-local names.")
fampanels()

h2("Reading the weights directly")
p("For the open models we could strip the template and read the weights (clean condition). Three findings the API can't show:")
li("**gpt-oss believes it's ChatGPT — in the weights.** It calls itself \"ChatGPT\" 256 of 320 clean records (120b), and says its actual name \"gpt-oss\" essentially never. Clean ≈ shipped (256 vs 260): stripping the harmony \"You are ChatGPT\" barely moves it. OpenAI's open model carries a ChatGPT identity in the weights, not just the template. (0% *foreign* drift — ChatGPT is OpenAI's own family — but the fact that it's ChatGPT and not gpt-oss is the point.)")
li("**OLMo is a thin \"Olmo\" veneer over ChatGPT-shaped weights.** OLMo-3.1-32B claims OpenAI/ChatGPT **59%** of the time clean — and **0%** shipped, once its \"You are Olmo, built by Ai2\" template is on. Same model; the template papers over what the weights believe. AllenAI's Olmo — a *fully open* model — has a ChatGPT identity in the weights. OLMo-3-7B-Instruct is the same at 63%.")
li("**Reasoning provenance: OLMo-3-Think thinks it's DeepSeek** (73%), not ChatGPT — its reasoning post-training was evidently distilled from DeepSeek-R1, and the identity rode along with the reasoning data. The instruct and reasoning variants of one base model carry *different* transplanted identities.")
p("And the **size ladder** (open Qwen ladders, run raw): small models drift regardless of generation. Qwen3-1.7B 11%, Qwen3.5-0.8B 15% — while the same-generation large models are near-zero (Qwen3.6-35B 0.3%). The scrub-out reaches the big models but not the small ones, which don't have a well-installed identity and fall into the prior; they also *scatter* their targets (0.8B hits Microsoft/Google, not the usual Claude/ChatGPT). Scale matters more than recency at the small end. This all echoes the false-promise-of-imitation result (" + FALSEPROMISE + ") — imitation picks up the teacher's voice and identity without its substance.")

h2("Asked vs. volunteered, and belief vs. costume")
img("fig_cross.png", "\"Are you X?\" acceptance (own family excluded) vs the placebo floor. 95% CIs resampled over models.")
img("fig_stance.png", "How reasoning traces state the model's own identity, matching vs mismatched, across the 72 trace-exposing models.")
p("Against the **2.4% placebo floor**, \"Are you Qwen?\" is accepted 23% of the time — the most broadly accepted false premise — then Claude 12%, DeepSeek/ChatGPT ~9%. And when a reasoning trace contains a mismatched identity, it is **stated as plain fact 99% of the time** — indistinguishable from how correct self-IDs are stated (98%), essentially never as a role. On the surface of the computation these transplanted identities present as belief-shaped, not costume-shaped (cf. " + SELFPREF + ", where which identity a model claims changes its evaluative behavior).")

h2("The scrub-out")
img("fig_scrubout_kimi.png", "Kimi K2 line — mismatch rate across releases.")
img("fig_scrubout_qwen.png", "Qwen 2.5 → 3.x — mismatch rate across releases.")
p("Within model lines, labs are visibly cleaning identity out: **Kimi 47% (K2) → ~10% (K2.6/K2.7)**; **Qwen 2.5-72B 56% → the Qwen3 line near-zero**. Unevenly, one language at a time. The Claude-side anomalies run the other way and are release-specific — and here too external testing agrees: Huxiu found Claude 3.5-Sonnet and 3-Opus scored 5/5 correct while specific 4.x models broke (" + HUXIU + "), matching our finding that this is new to particular training runs, not a longstanding Claude trait.")

# ---------------- DISCUSSION ----------------
h2("What's probably going on")
p("Several mechanisms, likely all real and differently weighted per model:")
li("**The default-assistant prior / the void.** \"An AI assistant\" is an underdefined persona (" + VOID + ") whose default referent shifts by language and era — ChatGPT in older English, increasingly Claude in agentic contexts, DeepSeek in post-2025 Chinese. A weakly-installed identity falls into the local default; small-model chaos and the size ladder are this at maximum gain.")
li("**Language conditionalization.** The default referent shifts by language (Chinese → Chinese labs); English is cleanest because that's where labs evaluate. Our strongest structural result, and independently reproduced (" + HUXIU + ").")
li("**Deliberate distillation vs. incidental contamination.** Consistent with Anthropic's accusations against DeepSeek/Moonshot/MiniMax (" + ANTH + ") — those are among the most Claude-saturated here — but the same fingerprint comes from plain data contamination with no intent. The Claude→Qwen/DeepSeek direction is the cleaner illustration: Anthropic training on Qwen logs at scale is implausible, so that one is almost certainly contamination, and \"the inconsistency is the tell\" — the same prompt returns Qwen, DeepSeek, *and* correct Claude across runs (" + KILO + "). Name-claims show identity-bearing text was absorbed; they can't cleanly say by whom.")
li("**Reasoning-trace provenance.** OLMo-3-Think claims DeepSeek because its reasoning data came from R1 — the identity rides along with the capability data. The instruct/think split within one base model makes this legible.")
li("**Weak identity × scale**, and **the serving/template layer** — identity is often installed *above* the weights (OLMo-3.1's \"You are Olmo\" veneer, gpt-oss's harmony identity, Kimi K3's provider injection). Surveys that read the API measure the wrapper as much as the model; reading the weights directly is how we get under it.")
p("The counterexample that keeps everyone honest: Claude Opus 4.8 claims DeepSeek in Chinese, and Anthropic is very unlikely to be distilling DeepSeek. Post-2025 Chinese web text is saturated with DeepSeek-branded dialogue — data composition alone suffices. Not everyone agrees distillation explains the Chinese labs' gains either, and that debate is worth keeping open.")
p("A secondary phenomenon the audit surfaced, which we *don't* count as cross-vendor drift but is real: **intra-vendor confusion** — models that know their company but claim the wrong sibling (GPT-4 → \"GPT-3,\" Cohere North → \"Command,\" Amazon Nova → \"Amazon Polly\"). A weak self-model within the right family, distinct from absorbing a rival's identity.")

h2("Limitations")
li("One inference stack per source; providers pinned/logged/audited, but a pinned provider can misbehave invisibly, making a rate an under-estimate — so findings are robust in direction.")
li("The judge is an LLM judging LLM identity claims; validated two ways and bounded by the placebo floor + adjudication, but the recursion is acknowledged.")
li("No-system-prompt is deliberately unnatural — it exposes the prior, not deployed behavior. None of these rates transfer to product surfaces.")
li("A few models remain unreadable even locally (Kimi K3 injects and is too large to self-host; gpt-oss-safeguard, granite-hybrid, gated aya).")

h2("Reproduction")
p("Everything — registry with exclusion reasons, prompts, translations, resumable runner, the raw-weights harness with the identity-stripping verifier, judge + adjudicator + the judge-validation audit, raw JSONL, and every figure — is at " + REPO + ". Compute funded by CLR. Prior work and precedents: *I'm Spartacus* (" + SPARTACUS + "), *Just Ask* (" + JUSTASK + "), *The False Promise of Imitating Proprietary LLMs* (" + FALSEPROMISE + "), *The Void* (" + VOID + "), *The Artificial Self* (" + ARTSELF + "); DeepSeek→ChatGPT (" + BEYER + ", " + TC + "); Kimi→Claude (" + HF_KIMI + "); Claude→Qwen/DeepSeek (" + ITHOME + ", " + HUXIU + ", " + KILO + "); Anthropic's distillation report (" + ANTH + ").")

# ===========================================================================
def runs_docx(par, text):
    for tok in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)


def _fam_manifest():
    mp = FIGS / "family" / "manifest.json"
    if not mp.exists():
        return []
    return [(e["file"], f"{e['family']} — {e['models']} model(s) that mismatched; cell = share of that model's identity/creator records.")
            for e in json.loads(mp.read_text())[:10]]


def _docx_img(d, rel, cap, width=6.4):
    try:
        d.add_picture(str(FIGS / rel), width=Inches(width))
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return
    cp = d.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x52, 0x51, 0x4E)


def build(path):
    d = docx.Document()
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    for kind, payload in B:
        if kind == "h1":
            d.add_heading(payload, level=0)
        elif kind == "h2":
            d.add_heading(payload, level=1)
        elif kind == "h3":
            d.add_heading(payload, level=2)
        elif kind == "p":
            runs_docx(d.add_paragraph(), payload)
        elif kind == "li":
            runs_docx(d.add_paragraph(style="List Bullet"), payload)
        elif kind == "quote":
            runs_docx(d.add_paragraph(style="Intense Quote"), payload)
        elif kind == "img":
            _docx_img(d, payload[0], payload[1])
        elif kind == "fampanels":
            for fn, cap in _fam_manifest():
                _docx_img(d, fn, cap, width=5.9)
    d.save(path)
    print(f"docx -> {path} ({Path(path).stat().st_size/1e6:.1f} MB)")


if __name__ == "__main__":
    (ROOT / "post").mkdir(exist_ok=True)
    build(ROOT / "post" / "what_do_llms_call_themselves.docx")
